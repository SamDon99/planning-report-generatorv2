
import streamlit as st
import os
import openai
from docx import Document

openai.api_key = st.secrets["OPENAI_API_KEY"]

st.title("📄 MHL Report Generator")

st.sidebar.header("Client + Report Type")
client = st.sidebar.selectbox("Client", ["MHL"])
report_type = st.sidebar.selectbox("Report Type", ["Mobility Management Plan"])

input_text = st.text_area("Paste your project brief (e.g. Longview N7)", height=400)

if st.button("🛠 Generate Report"):
    with st.spinner("Generating report..."):
        prompt = f'''
You are writing a professional Mobility Management Plan for {client}.

Use this structure:
1. Introduction
2. Site Context
3. Access Strategy
4. Sustainable Transport
5. DMURS Compliance
6. Cycling & Walking
7. Parking Strategy
8. Public Transport Links
9. Conclusions

Target length: 1800–2500 words.
Match the tone of Irish planning reports.
Reference Cork City Development Plan and DMURS where relevant.

Project details:
{input_text}
'''

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )

        report_text = response["choices"][0]["message"]["content"]

        doc = Document()
        doc.add_heading("Mobility Management Plan", 0)
        doc.add_paragraph(report_text)

        os.makedirs("outputs", exist_ok=True)
        output_path = "outputs/MHL_Mobility_Management_Plan.docx"
        doc.save(output_path)

        st.success("✅ Report ready!")
        with open(output_path, "rb") as f:
            st.download_button("📥 Download Report", f, file_name="MHL_Mobility_Management_Plan.docx")
